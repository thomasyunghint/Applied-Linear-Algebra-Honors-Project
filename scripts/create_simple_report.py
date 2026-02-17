"""Create a simple, easy-to-understand 1-page Word document for the professor."""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main() -> None:
    doc = Document()
    
    # Title
    title = doc.add_heading("MAT 343 Honors Contract Project Summary", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project Overview
    doc.add_heading("Project Overview", level=2)
    doc.add_paragraph(
        "For this project, I wanted to see how linear algebra actually works in finance. I ended up "
        "focusing on three main areas: analyzing option price sensitivity using the Greeks (Delta, Gamma, "
        "Theta), using systems of equations to hedge a portfolio and reduce risk by about 95%, and "
        "optimizing a portfolio of four stocks—GME, ZM, TSLA, and MCD—to find the minimum-variance "
        "combination."
    )
    
    # Key Findings
    doc.add_heading("Key Findings", level=2)
    doc.add_paragraph(
        "The option analysis was interesting—I found that sensitivity peaks near the strike prices, "
        "and Delta and Gamma act a lot like derivatives, which makes sense mathematically. For the "
        "hedging part, I set it up as Ax = b where the risk matrix times the hedge weights equals "
        "the target risk, and this actually eliminated about 95% of the portfolio risk. The portfolio "
        "optimization results were pretty surprising: the minimum-variance portfolio put 74% into "
        "MCD (which is pretty stable) and only 5% into GME (which is super volatile). This gave me "
        "13% annual volatility instead of 21% if I just split everything equally."
    )
    
    # Linear Algebra Connections
    doc.add_heading("Linear Algebra Applications", level=2)
    
    doc.add_paragraph(
        "When I started working on the hedging problem, I realized it was basically a system of "
        "equations. I set it up as Ax = b, where A is the risk matrix—each row represents a different "
        "type of risk, and the columns are the different hedging instruments I could use. The x vector "
        "contains the hedge weights I'm trying to find. It was really satisfying to see how the methods "
        "we learned in class actually work for this real problem."
    )
    
    doc.add_paragraph(
        "For the portfolio optimization, I used the formula w = (Σ⁻¹1)/(1ᵀΣ⁻¹1) where Σ is the covariance "
        "matrix. At first I didn't really understand why inverting the covariance matrix would give me "
        "the optimal weights, but after working through it, I realized the inversion captures how the "
        "stocks interact with each other. It's not just a math trick—the structure of the inverse "
        "matrix actually tells you which stocks to weight more heavily. The fact that this gave me "
        "13% volatility instead of 21% really shows how powerful matrix inversion can be for optimization."
    )
    
    doc.add_paragraph(
        "The quadratic form wᵀΣw is what calculates portfolio risk, and honestly this turned out to be "
        "the foundation of everything I did. When I plotted the efficient frontier, I started to see "
        "that each point on that curve represents a different portfolio, and they're all just variations "
        "of this same matrix multiplication. It's kind of amazing that one operation can capture the "
        "whole risk-return relationship."
    )
    
    doc.add_paragraph(
        "I also noticed some connections to calculus that we've covered. The Delta and Gamma of options "
        "are related to Jacobian and Hessian matrices. Delta is like a first derivative—it gives you a "
        "linear approximation of how the option price changes. When you add Gamma, you get a quadratic "
        "approximation, which is like the second derivative. It's cool how linear algebra gives you the "
        "framework to think about these approximations."
    )
    
    doc.add_paragraph(
        "The correlation matrices were probably the most intuitive part. When stocks have low correlation, "
        "they naturally diversify each other. I made some heatmaps to visualize this, and you can "
        "literally see how the symmetric structure of the matrix corresponds to diversification. The "
        "lower the correlations, the better the portfolio can reduce risk."
    )
    
    # Significance
    doc.add_heading("Why This Matters", level=2)
    doc.add_paragraph(
        "What I learned from this project is that linear algebra isn't just abstract theory—it actually "
        "produces real, measurable results. The fact that I could reduce risk by 95% through hedging "
        "shows that systems of equations can solve real problems. And the 38% volatility reduction "
        "(going from 21% to 13%) proves that mathematical optimization works way better than just "
        "splitting everything equally. This project really showed me that the concepts from MAT 343 "
        "are essential tools for quantitative finance, where matrix operations directly lead to better "
        "investment outcomes."
    )
    
    # Conclusion
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "Overall, this project showed me that linear algebra is a really practical tool for financial "
        "risk management. Every calculation I did—from finding hedge weights to optimizing the portfolio—"
        "relied on matrix operations and systems of equations. The results speak for themselves: "
        "mathematical optimization can improve investment outcomes significantly, reducing risk by 38% "
        "while still maintaining returns."
    )
    
    # Save
    output_path = Path(__file__).parent.parent / "reports" / "project_summary_simple.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Created simple report at {output_path}")


if __name__ == "__main__":
    main()
